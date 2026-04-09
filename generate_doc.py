from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

doc = Document()

# ============================================================================
# PAGE SETUP
# ============================================================================
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ============================================================================
# STYLE SETUP
# ============================================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Heading styles
for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_style.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    heading_style.font.name = 'Calibri'
    if i == 1:
        heading_style.font.size = Pt(24)
        heading_style.font.bold = True
    elif i == 2:
        heading_style.font.size = Pt(18)
        heading_style.font.bold = True
    elif i == 3:
        heading_style.font.size = Pt(14)
        heading_style.font.bold = True

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================
def set_cell_shading(cell, color_hex):
    """Set cell background color"""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def add_formatted_table(headers, rows, header_color="1F497D"):
    """Add a formatted table with header styling"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)
        set_cell_shading(cell, header_color)
    
    # Data rows
    for row_idx, row_data in enumerate(rows, 1):
        for col_idx, cell_data in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = str(cell_data)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F2F2F2")
    
    doc.add_paragraph()
    return table

def add_bullet_point(text, bold=False):
    """Add a bullet point"""
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.bold = bold
    return p

# ============================================================================
# COVER PAGE
# ============================================================================
doc.add_paragraph()
doc.add_paragraph()

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('DOKUMENTASI TEKNIS')
run.font.size = Pt(36)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('MCP PostgreSQL Chatbot & ERP Integration')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

doc.add_paragraph()

# System info box
info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

info_data = [
    ('Nama Aplikasi', 'MCP PostgreSQL Chatbot & ERP Integration'),
    ('Versi', '1.0.0'),
    ('Framework', 'Laravel 12.x'),
    ('Database', 'PostgreSQL 12+'),
    ('Bahasa Pemrograman', 'PHP 8.2+'),
    ('Tanggal Dokumentasi', datetime.datetime.now().strftime('%d %B %Y'))
]

for i, (label, value) in enumerate(info_data):
    cell_label = info_table.rows[i].cells[0]
    cell_value = info_table.rows[i].cells[1]
    
    cell_label.text = label
    cell_value.text = value
    
    for cell in [cell_label, cell_value]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(11)
    
    cell_label.paragraphs[0].runs[0].font.bold = True
    set_cell_shading(cell_label, "E8F0FE")
    cell_label.width = Cm(5)
    cell_value.width = Cm(10)

doc.add_paragraph()
doc.add_paragraph()

# Footer info
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run('Dokumen ini berisi dokumentasi lengkap sistem\nTermasuk arsitektur, instalasi, konfigurasi, dan panduan operasional')
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

doc.add_page_break()

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
doc.add_heading('DAFTAR ISI', level=1)

toc_items = [
    '1. Gambaran Umum Sistem',
    '   1.1 Deskripsi Aplikasi',
    '   1.2 Fitur Utama',
    '   1.3 Arsitektur Sistem',
    '2. Persyaratan Sistem',
    '   2.1 Server Requirements',
    '   2.2 Dependencies',
    '3. Instalasi dan Setup',
    '   3.1 Instalasi di Windows',
    '   3.2 Instalasi di Linux/Ubuntu',
    '   3.3 Konfigurasi Database',
    '4. Struktur Database',
    '   4.1 Diagram Relasi Tabel',
    '   4.2 Detail Tabel',
    '5. Routing dan API Endpoints',
    '   5.1 Public Routes',
    '   5.2 Protected Routes',
    '   5.3 Admin Routes',
    '6. Fitur Chatbot AI',
    '   6.1 Agentic Tool Calling',
    '   6.2 Integrasi ERP',
    '   6.3 Export Data',
    '7. Role-Based Access Control (RBAC)',
    '   7.1 Konsep RBAC',
    '   7.2 Manajemen Role',
    '8. Admin Panel',
    '   8.1 Manajemen User',
    '   8.2 Manajemen Role',
    '   8.3 Import/Export',
    '9. MCP Server Integration',
    '10. Optimasi Performa',
    '    10.1 Database Indexing',
    '    10.2 Caching Strategy',
    '    10.3 Laravel Octane',
    '11. Deployment',
    '    11.1 Deployment di Ubuntu',
    '    11.2 Production Checklist',
    '12. Troubleshooting',
    '13. Maintenance dan Monitoring',
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0
    if not item.startswith('   '):
        for run in p.runs:
            run.font.bold = True

doc.add_page_break()

# ============================================================================
# 1. GAMBARAN UMUM SISTEM
# ============================================================================
doc.add_heading('1. Gambaran Umum Sistem', level=1)

doc.add_heading('1.1 Deskripsi Aplikasi', level=2)
doc.add_paragraph(
    'MCP PostgreSQL Chatbot & ERP Integration adalah aplikasi web berbasis Laravel 12.x yang '
    'mengintegrasikan chatbot cerdas dengan database PostgreSQL dan dokumentasi ERP. Aplikasi ini '
    'dirancang untuk membantu analisis data bisnis dan memberikan panduan operasional ERP secara '
    'interaktif menggunakan kecerdasan buatan (Artificial Intelligence).'
)

doc.add_paragraph(
    'Sistem ini memanfaatkan Model Context Protocol (MCP) untuk menyediakan kemampuan AI yang '
    'dapat memahami konteks bisnis, menjalankan query database, dan memberikan insight yang '
    'relevan kepada pengguna sesuai dengan role dan hak akses masing-masing.'
)

doc.add_heading('1.2 Fitur Utama', level=2)

features = [
    ('Analisis Data Bisnis', 
     'Terhubung langsung ke database PostgreSQL untuk menjawab pertanyaan seputar penjualan, '
     'produk terlaris, retensi pelanggan, dan metrik bisnis lainnya secara real-time.'),
    ('Panduan ERP (ERP Guidance)', 
     'Mengambil data dari erp-guidance.online dan memberikan instruksi langkah demi langkah '
     'untuk operasional ERP, termasuk order pembelian, permintaan pembelian, dan klaim barang.'),
    ('Premium UI dengan Glassmorphism', 
     'Tampilan modern dengan tema Glassmorphism yang memberikan pengalaman pengguna yang '
     'menarik dan intuitif.'),
    ('User Authentication', 
     'Sistem login dan register yang aman untuk melindungi akses ke aplikasi dan data bisnis.'),
    ('Role-Based Access Control (RBAC)', 
     'Pembatasan akses tabel database berdasarkan role user, memastikan setiap pengguna '
     'hanya dapat mengakses data yang menjadi haknya.'),
    ('Admin Dashboard', 
     'Panel administrasi untuk manajemen user, role, dan permissions melalui UI drag & drop '
     'yang mudah digunakan.'),
    ('MCP Server Integration', 
     'Integrasi dengan Model Context Protocol untuk AI assistance dan tool calling yang '
     'memungkinkan AI berinteraksi dengan database dan sistem ERP.'),
    ('Export Data (Excel & PDF)', 
     'Kemampuan untuk export data chat dan analisis ke format Excel dan PDF untuk '
     'keperluan pelaporan dan dokumentasi.'),
]

for title, desc in features:
    p = doc.add_paragraph()
    run = p.add_run(f'✓ {title}: ')
    run.font.bold = True
    p.add_run(desc)

doc.add_heading('1.3 Arsitektur Sistem', level=2)

doc.add_paragraph(
    'Aplikasi ini menggunakan arsitektur Model-View-Controller (MVC) dengan struktur sebagai berikut:'
)

architecture_items = [
    'Frontend: Blade templates dengan Tailwind CSS dan JavaScript untuk interaktivitas',
    'Backend: Laravel 12.x dengan PHP 8.2+',
    'Database: PostgreSQL 12+ untuk primary database',
    'Cache: Redis untuk caching dan session management',
    'AI Integration: OpenAI API (GPT models) dengan fallback mechanism',
    'Queue System: Database/Redis queue untuk background jobs',
    'Server: Nginx/Apache dengan PHP-FPM atau Laravel Octane (Swoole)'
]

for item in architecture_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 2. PERSYARATAN SISTEM
# ============================================================================
doc.add_heading('2. Persyaratan Sistem', level=1)

doc.add_heading('2.1 Server Requirements', level=2)

req_table = [
    ['Komponen', 'Versi Minimum', 'Versi Rekomendasi'],
    ['PHP', '8.2', '8.2+'],
    ['PostgreSQL', '12', '14+'],
    ['Composer', '2.0+', 'Latest'],
    ['Node.js', '18.x', '20.x LTS'],
    ['NPM', '8.x', 'Latest'],
    ['Redis', '6.x', '7.x'],
    ['Nginx', '1.18+', '1.24+'],
    ['RAM (Minimum)', '2 GB', '4+ GB'],
    ['Storage', '10 GB', '20+ GB SSD'],
]

add_formatted_table(req_table[0], req_table[1:])

doc.add_heading('2.2 Dependencies', level=2)

doc.add_heading('PHP Dependencies (composer.json)', level=3)

php_deps = [
    'laravel/framework ^12.0 - Framework utama',
    'php-mcp/server - MCP Server implementation',
    'php-mcp/laravel - MCP integration untuk Laravel',
    'symfony/css-selector - DOM parsing',
    'symfony/dom-crawler - Web crawling untuk ERP documentation',
    'maatwebsite/excel - Export/import Excel',
    'barryvdh/laravel-dompdf - PDF generation',
]

for dep in php_deps:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_heading('JavaScript Dependencies (package.json)', level=3)

js_deps = [
    'tailwindcss ^4.0.0 - CSS framework',
    'vite ^7.0.7 - Build tool',
    'axios ^1.11.0 - HTTP client',
    'cheerio ^1.2.0 - Server-side DOM parsing',
    'concurrently ^9.0.1 - Multi-process runner',
]

for dep in js_deps:
    doc.add_paragraph(dep, style='List Bullet')

doc.add_heading('PHP Extensions Required', level=3)

extensions = [
    'php8.2-pgsql - PostgreSQL driver',
    'php8.2-mbstring - Multibyte string support',
    'php8.2-xml - XML parsing',
    'php8.2-curl - HTTP requests',
    'php8.2-zip - ZIP archive support',
    'php8.2-redis - Redis client (optional, recommended)',
    'php8.2-swoole - Swoole extension untuk Octane (optional)',
]

for ext in extensions:
    doc.add_paragraph(ext, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 3. INSTALASI DAN SETUP
# ============================================================================
doc.add_heading('3. Instalasi dan Setup', level=1)

doc.add_heading('3.1 Instalasi di Windows', level=2)

steps = [
    ('Clone Repository', [
        'cd "D:\\MCP Versi Web"',
        'git clone <repository-url> mcp-postgresql',
        'cd mcp-postgresql'
    ]),
    ('Instalasi Dependensi', [
        'composer install',
        'npm install'
    ]),
    ('Konfigurasi Environment', [
        'copy .env.example .env',
        'php artisan key:generate'
    ]),
    ('Konfigurasi Database', [
        'Edit file .env dan sesuaikan konfigurasi database:',
        'DB_CONNECTION=pgsql',
        'DB_HOST=127.0.0.1',
        'DB_PORT=5432',
        'DB_DATABASE=db_penjualan',
        'DB_USERNAME=postgres',
        'DB_PASSWORD=your_password'
    ]),
    ('Migrasi Database', [
        'php artisan migrate'
    ]),
    ('Seed Data Default', [
        'php artisan db:seed'
    ]),
    ('Build Frontend', [
        'npm run build'
    ]),
    ('Jalankan Development Server', [
        'php artisan serve',
        'Akses aplikasi di: http://localhost:8000'
    ]),
]

for i, (step, commands) in enumerate(steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {step}')
    run.font.bold = True
    
    for cmd in commands:
        doc.add_paragraph(cmd).style = 'No Spacing'
    doc.add_paragraph()

doc.add_heading('3.2 Instalasi di Linux/Ubuntu', level=2)

linux_steps = [
    ('Instalasi Dependensi Sistem', [
        'sudo apt update',
        'sudo apt install php8.2 php8.2-cli php8.2-pgsql php8.2-mbstring \\',
        '  php8.2-xml php8.2-curl php8.2-zip php8.2-gd unzip git curl nodejs npm -y'
    ]),
    ('Instalasi Composer', [
        'curl -sS https://getcomposer.org/installer | php',
        'sudo mv composer.phar /usr/local/bin/composer'
    ]),
    ('Clone Repository', [
        'cd /var/www',
        'git clone <repository-url> mcp-postgresql',
        'cd mcp-postgresql'
    ]),
    ('Instalasi Dependensi', [
        'composer install',
        'npm install',
        'npm run build'
    ]),
    ('Konfigurasi Environment', [
        'cp .env.example .env',
        'php artisan key:generate'
    ]),
    ('Migrasi & Seed', [
        'php artisan migrate --seed'
    ]),
    ('Set Permissions', [
        'sudo chown -R www-data:www-data /var/www/mcp-postgresql',
        'sudo chmod -R 755 /var/www/mcp-postgresql/storage'
    ]),
]

for i, (step, commands) in enumerate(linux_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {step}')
    run.font.bold = True
    
    for cmd in commands:
        doc.add_paragraph(cmd).style = 'No Spacing'
    doc.add_paragraph()

doc.add_heading('3.3 Konfigurasi Database', level=2)

doc.add_paragraph('Buat database PostgreSQL:')

db_commands = [
    'sudo -u postgres psql',
    '',
    'CREATE DATABASE db_penjualan;',
    'CREATE USER postgres WITH PASSWORD \'your_password\';',
    'GRANT ALL PRIVILEGES ON DATABASE db_penjualan TO postgres;',
    '\\q'
]

for cmd in db_commands:
    doc.add_paragraph(cmd).style = 'No Spacing'

doc.add_page_break()

# ============================================================================
# 4. STRUKTUR DATABASE
# ============================================================================
doc.add_heading('4. Struktur Database', level=1)

doc.add_heading('4.1 Diagram Relasi Tabel', level=2)

doc.add_paragraph(
    'Database terdiri dari tabel-tabel berikut yang saling berelasi untuk mendukung '
    'fungsionalitas aplikasi:'
)

tables_overview = [
    ['Tabel', 'Fungsi', 'Jumlah Field'],
    ['users', 'Data pengguna aplikasi', '8'],
    ['roles', 'Definisi role/hak akses', '4'],
    ['role_permissions', 'Permissions setiap role', '4'],
    ['chat_sessions', 'Sesi percakapan chatbot', '4'],
    ['chat_messages', 'Pesan dalam setiap sesi chat', '6'],
    ['documentation', 'Dokumentasi ERP yang di-crawl', '6'],
    ['sessions', 'Session management (Laravel)', '6'],
    ['cache', 'Cache storage', '4'],
    ['jobs', 'Queue jobs', '7'],
    ['failed_jobs', 'Failed jobs tracking', '6'],
]

add_formatted_table(tables_overview[0], tables_overview[1:])

doc.add_heading('4.2 Detail Tabel', level=2)

# Users table
doc.add_heading('4.2.1 Tabel users', level=3)
doc.add_paragraph('Menyimpan data pengguna aplikasi termasuk informasi login dan role.')

users_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['name', 'VARCHAR', 'NOT NULL', 'Nama lengkap user'],
    ['email', 'VARCHAR', 'UNIQUE, NOT NULL', 'Email user'],
    ['email_verified_at', 'TIMESTAMP', 'NULLABLE', 'Waktu verifikasi email'],
    ['password', 'VARCHAR', 'NOT NULL', 'Password (hashed)'],
    ['role', 'BIGINT', 'FOREIGN KEY', 'Referensi ke tabel roles'],
    ['is_admin', 'BOOLEAN', 'DEFAULT false', 'Flag user sebagai admin'],
    ['remember_token', 'VARCHAR', 'NULLABLE', 'Token remember me'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(users_fields[0], users_fields[1:])

# Roles table
doc.add_heading('4.2.2 Tabel roles', level=3)
doc.add_paragraph('Menyimpan definisi role untuk sistem RBAC.')

roles_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['name', 'VARCHAR', 'UNIQUE, NOT NULL', 'Nama role'],
    ['description', 'VARCHAR', 'NULLABLE', 'Deskripsi role'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(roles_fields[0], roles_fields[1:])

# Role permissions table
doc.add_heading('4.2.3 Tabel role_permissions', level=3)
doc.add_paragraph('Menyimpan mapping permission (akses tabel) untuk setiap role.')

perms_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['role_id', 'BIGINT', 'FOREIGN KEY', 'Referensi ke tabel roles'],
    ['table_name', 'VARCHAR', 'NOT NULL', 'Nama tabel yang dapat diakses'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(perms_fields[0], perms_fields[1:])

doc.add_paragraph(
    'Catatan: Terdapat unique constraint pada kombinasi [role_id, table_name] untuk '
    'mencegah duplikasi permission.'
)

# Chat sessions table
doc.add_heading('4.2.4 Tabel chat_sessions', level=3)
doc.add_paragraph('Menyimpan sesi percakapan setiap user dengan chatbot.')

sessions_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['user_id', 'BIGINT', 'FOREIGN KEY', 'Referensi ke tabel users'],
    ['title', 'VARCHAR', 'NOT NULL', 'Judul sesi chat'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(sessions_fields[0], sessions_fields[1:])

doc.add_paragraph(
    'Relasi: Satu user dapat memiliki banyak chat session (one-to-many). '
    'Jika user dihapus, semua session terkait akan dihapus (cascade).'
)

# Chat messages table
doc.add_heading('4.2.5 Tabel chat_messages', level=3)
doc.add_paragraph('Menyimpan pesan-pesan dalam setiap sesi percakapan.')

messages_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['chat_session_id', 'BIGINT', 'FOREIGN KEY', 'Referensi ke tabel chat_sessions'],
    ['role', 'VARCHAR', 'NOT NULL', 'Role pesan (user/assistant/tool)'],
    ['content', 'TEXT', 'NULLABLE', 'Konten pesan'],
    ['tool_results', 'TEXT', 'NULLABLE, JSON', 'Hasil eksekusi tool (jika ada)'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(messages_fields[0], messages_fields[1:])

doc.add_paragraph(
    'Relasi: Satu session dapat memiliki banyak messages (one-to-many). '
    'Jika session dihapus, semua messages terkait akan dihapus (cascade).'
)

# Documentation table
doc.add_heading('4.2.6 Tabel documentation', level=3)
doc.add_paragraph('Menyimpan dokumentasi ERP yang diambil dari website erp-guidance.online.')

doc_fields = [
    ['Field', 'Tipe', 'Constraint', 'Deskripsi'],
    ['id', 'BIGINT', 'PRIMARY KEY', 'Auto-increment ID'],
    ['title', 'VARCHAR', 'NOT NULL', 'Judul dokumentasi'],
    ['url', 'VARCHAR', 'UNIQUE, NOT NULL', 'URL sumber dokumentasi'],
    ['content', 'TEXT', 'NOT NULL', 'Konten dokumentasi'],
    ['category', 'VARCHAR', 'NULLABLE', 'Kategori dokumentasi'],
    ['created_at', 'TIMESTAMP', 'NULLABLE', 'Waktu pembuatan record'],
    ['updated_at', 'TIMESTAMP', 'NULLABLE', 'Waktu update terakhir'],
]

add_formatted_table(doc_fields[0], doc_fields[1:])

doc.add_page_break()

# ============================================================================
# 5. ROUTING DAN API ENDPOINTS
# ============================================================================
doc.add_heading('5. Routing dan API Endpoints', level=1)

doc.add_heading('5.1 Public Routes', level=2)

public_routes = [
    ['Method', 'URL', 'Controller', 'Deskripsi'],
    ['GET', '/login', 'AuthController@showLogin', 'Halaman login'],
    ['POST', '/login', 'AuthController@login', 'Proses autentikasi'],
    ['GET', '/register', 'AuthController@showRegister', 'Halaman register (local only)'],
    ['POST', '/register', 'AuthController@register', 'Proses registrasi (local only)'],
    ['POST', '/logout', 'AuthController@logout', 'Logout user'],
    ['GET', '/mcp', 'MCPServer', 'MCP Server endpoint'],
]

add_formatted_table(public_routes[0], public_routes[1:])

doc.add_heading('5.2 Protected Routes (Authenticated)', level=2)

protected_routes = [
    ['Method', 'URL', 'Controller', 'Deskripsi'],
    ['GET', '/', 'Redirect', 'Redirect ke /chatbot'],
    ['GET', '/chatbot', 'AgenticChatbotController@index', 'Halaman utama chatbot'],
    ['POST', '/chatbot/send', 'AgenticChatbotController@send', 'Kirim pesan ke chatbot'],
    ['POST', '/chatbot/export/excel', 'AgenticChatbotController@exportExcel', 'Export chat ke Excel'],
    ['POST', '/chatbot/export/pdf', 'AgenticChatbotController@exportPdf', 'Export chat ke PDF'],
    ['GET', '/chatbot/sessions', 'AgenticChatbotController@getSessions', 'Daftar sesi chat user'],
    ['GET', '/chatbot/sessions/{id}', 'AgenticChatbotController@getSession', 'Detail sesi chat'],
    ['DELETE', '/chatbot/sessions/{id}', 'AgenticChatbotController@deleteSession', 'Hapus sesi chat'],
    ['PUT', '/chatbot/sessions/{id}', 'AgenticChatbotController@updateSessionTitle', 'Update judul sesi'],
]

add_formatted_table(protected_routes[0], protected_routes[1:])

doc.add_heading('5.3 Admin Routes', level=2)

admin_routes = [
    ['Method', 'URL', 'Controller', 'Deskripsi'],
    ['GET', '/admin', 'AdminController@index', 'Dashboard admin'],
    ['GET', '/admin/users', 'AdminController@users', 'Manajemen user'],
    ['POST', '/admin/users', 'AdminController@userStore', 'Tambah user'],
    ['PUT', '/admin/users/{user}', 'AdminController@userUpdate', 'Update user'],
    ['DELETE', '/admin/users/{user}', 'AdminController@userDelete', 'Hapus user'],
    ['GET', '/admin/users/export', 'AdminController@usersExport', 'Export users ke Excel'],
    ['POST', '/admin/users/import', 'AdminController@usersImport', 'Import users dari Excel'],
    ['GET', '/admin/users/template', 'AdminController@userTemplate', 'Download template Excel'],
    ['GET', '/admin/roles', 'AdminController@roles', 'Manajemen role'],
    ['POST', '/admin/roles', 'AdminController@roleStore', 'Tambah role'],
    ['PUT', '/admin/roles/{role}', 'AdminController@roleUpdate', 'Update role'],
    ['DELETE', '/admin/roles/{role}', 'AdminController@roleDelete', 'Hapus role'],
    ['POST', '/admin/roles/{role}/permissions', 'AdminController@updatePermissions', 'Update permissions'],
]

add_formatted_table(admin_routes[0], admin_routes[1:])

doc.add_page_break()

# ============================================================================
# 6. FITUR CHATBOT AI
# ============================================================================
doc.add_heading('6. Fitur Chatbot AI', level=1)

doc.add_heading('6.1 Agentic Tool Calling', level=2)

doc.add_paragraph(
    'Chatbot menggunakan pendekatan Agentic Tool Calling, di mana AI tidak hanya memberikan '
    'respons tekstual, tetapi juga dapat menjalankan berbagai tools untuk mengakses data, '
    'menganalisis informasi, dan memberikan insight yang lebih mendalam.'
)

doc.add_heading('Model AI yang Digunakan', level=3)

models = [
    ['Prioritas', 'Model', 'Penggunaan'],
    ['Primary', 'gpt-5.4', 'Model utama untuk respons terbaik'],
    ['Fallback 1', 'gpt-5.4-mini', 'Fallback jika model utama rate limit'],
    ['Fallback 2', 'gpt-5.4-nano', 'Fallback kedua untuk ketersediaan'],
    ['Fallback 3', 'gpt-5.4-pro', 'Fallback dengan kemampuan advanced'],
]

add_formatted_table(models[0], models[1:])

doc.add_heading('Tool yang Tersedia', level=3)

tools = [
    ['Nama Tool', 'Fungsi', 'Deskripsi'],
    ['execute_query', 'Eksekusi SQL Query', 'Menjalankan query SQL ke database PostgreSQL dengan RBAC'],
    ['get_schema_info', 'Informasi Schema', 'Mendapatkan informasi struktur tabel dan relasi'],
    ['describe_table', 'Deskripsi Tabel', 'Menampilkan detail kolom dan data dalam tabel'],
    ['get_relationships', 'Relasi Tabel', 'Mendapatkan informasi relasi antar tabel'],
    ['get_indexes', 'Informasi Index', 'Menampilkan index yang ada pada tabel'],
    ['data_quality_check', 'Cek Kualitas Data', 'Memvalidasi kualitas dan konsistensi data'],
    ['statistical_analysis', 'Analisis Statistik', 'Menjalankan analisis statistik (13 metode)'],
    ['smart_analyze', 'Analisis Cerdas', 'Analisis query dengan rekomendasi otomatis'],
    ['explain_query_plan', 'Explain Query', 'Menampilkan execution plan query SQL'],
    ['query_templates', 'Template Query', 'Template query untuk analisis umum'],
    ['erp_navigation', 'Navigasi ERP', 'Navigasi menu ERP guidance'],
    ['erp_guidance', 'Panduan ERP', 'Memberikan instruksi operasional ERP'],
    ['language_detection', 'Deteksi Bahasa', 'Mendeteksi bahasa input user'],
]

add_formatted_table(tools[0], tools[1:])

doc.add_heading('Konfigurasi Tool Calling', level=3)

config_items = [
    'Max Tool Loops: 20 iterasi per permintaan',
    'Max History: 20 pesan terakhir dalam context',
    'Max Tokens: 32,768 tokens untuk respons lengkap',
    'Language Detection: Otomatis deteksi bahasa user (Indonesia/English)',
    'RBAC Enforcement: Query dibatasi sesuai role user',
]

for item in config_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('6.2 Integrasi ERP', level=2)

doc.add_paragraph(
    'Sistem terintegrasi dengan ERP Guidance (erp-guidance.online) untuk memberikan panduan '
    'operasional ERP kepada user. Integrasi dilakukan melalui web crawling dan penyimpanan '
    'dokumentasi di database lokal.'
)

doc.add_heading('Proses Crawling Dokumentasi ERP', level=3)

crawl_steps = [
    'Login ke ERP Guidance menggunakan credentials yang dikonfigurasi',
    'Mengambil daftar semua halaman dokumentasi dari sitemap',
    'Meng-crawl setiap halaman dan menyimpan konten (teks, gambar, video)',
    'Menyimpan data ke tabel documentation',
    'Enrichment dengan detail field formulir untuk halaman spesifik',
]

for i, step in enumerate(crawl_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Perintah Artisan untuk ERP Documentation', level=3)

erp_commands = [
    ['Perintah', 'Fungsi'],
    ['php artisan app:crawl-documentation', 'Mengambil semua dokumentasi dari website ERP'],
    ['php artisan app:enrich-documentation', 'Menambahkan detail field formulir ke dokumentasi'],
]

add_formatted_table(erp_commands[0], erp_commands[1:])

doc.add_paragraph(
    'Catatan Penting: Data dokumentasi ERP tidak disimpan di Git. Setiap server baru harus '
    'menjalankan crawler ini setelah deployment.'
)

doc.add_heading('Halaman ERP yang Di-enrich', level=3)

enriched_pages = [
    ['Halaman', 'Konten Tambahan'],
    ['Order Pembelian', 'Field header, detail barang, ringkasan total'],
    ['Permintaan Pembelian', 'Field header, tabel detail barang, tombol aksi'],
    ['Klaim Barang', 'Field header, detail barang/TTB, faktur pembelian, nilai klaim'],
]

add_formatted_table(enriched_pages[0], enriched_pages[1:])

doc.add_heading('6.3 Export Data', level=2)

doc.add_paragraph(
    'Chatbot menyediakan fitur export data percakapan dan hasil analisis ke berbagai format:'
)

export_formats = [
    ['Format', 'Library', 'Penggunaan'],
    ['Excel (.xlsx)', 'Maatwebsite Excel', 'Export dengan formatting dan chart'],
    ['PDF', 'DomPDF', 'Export untuk laporan dan dokumentasi'],
]

add_formatted_table(export_formats[0], export_formats[1:])

doc.add_paragraph(
    'Fitur export dapat diakses melalui tombol "Export to Excel" dan "Export to PDF" '
    'di halaman chatbot. Export mencakup:'
)

export_items = [
    'History percakapan (user messages dan assistant responses)',
    'Hasil query dan analisis yang dijalankan',
    'Informasi chart/data visualisasi (jika ada)',
    'Timestamp dan informasi sesi',
]

for item in export_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 7. ROLE-BASED ACCESS CONTROL (RBAC)
# ============================================================================
doc.add_heading('7. Role-Based Access Control (RBAC)', level=1)

doc.add_heading('7.1 Konsep RBAC', level=2)

doc.add_paragraph(
    'Sistem RBAC membatasi akses tabel database berdasarkan role user. Setiap role memiliki '
    'permissions yang menentukan tabel mana saja yang dapat diakses. Chatbot akan menjalankan '
    'query hanya pada tabel yang diizinkan untuk role user yang bersangkutan.'
)

doc.add_heading('Role Default', level=3)

roles_table = [
    ['Role ID', 'Nama Role', 'Deskripsi', 'Akses Tabel'],
    ['1', 'Data Entry', 'Akses transaksi & pembeli', 'transaksi, pembeli'],
    ['2', 'Produk Analyst', 'Akses transaksi & produk', 'transaksi, produk'],
    ['3', 'Super Admin', 'Akses seluruh tabel', 'Semua tabel'],
]

add_formatted_table(roles_table[0], roles_table[1:])

doc.add_heading('User Default (Seeder)', level=3)

users_table = [
    ['Email', 'Password', 'Role', 'Akses'],
    ['role1@example.com', 'password', 'Data Entry', 'transaksi, pembeli'],
    ['role2@example.com', 'password', 'Produk Analyst', 'transaksi, produk'],
    ['role3@example.com', 'password', 'Super Admin', 'Semua tabel'],
]

add_formatted_table(users_table[0], users_table[1:])

doc.add_paragraph(
    '⚠️ PENTING: Segera ubah password default setelah instalasi pertama kali!',
).runs[0].font.bold = True

doc.add_heading('7.2 Cara Kerja RBAC', level=2)

rbac_steps = [
    'User login dengan credentials',
    'Sistem mengecek role user dari database',
    'Sistem mengambil permissions (daftar tabel) untuk role tersebut',
    'Chatbot hanya dapat mengakses tabel yang ada dalam permissions role user',
    'Admin dapat mengubah permissions melalui panel administrasi (/admin/roles)',
    'Setiap query yang dijalankan divalidasi terhadap allowed tables',
]

for i, step in enumerate(rbac_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('Implementasi RBAC dalam Kode', level=3)

doc.add_paragraph(
    'RBAC diimplementasikan melalui beberapa komponen:'
)

rbac_components = [
    'Role Model: Menyimpan informasi role (nama, deskripsi)',
    'RolePermission Model: Mapping role ke tabel yang dapat diakses',
    'User Model: Memiliki relasi belongsTo ke Role',
    'ToolCallExecutor: Mengecek allowed tables sebelum menjalankan query',
    'QueryService: Menerima allowed tables dan memfilter query',
]

for comp in rbac_components:
    doc.add_paragraph(comp, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 8. ADMIN PANEL
# ============================================================================
doc.add_heading('8. Admin Panel', level=1)

doc.add_paragraph(
    'Admin panel menyediakan antarmuka grafis untuk manajemen user, role, dan permissions. '
    'Hanya user dengan is_admin = true yang dapat mengakses panel ini.'
)

doc.add_heading('8.1 Dashboard Admin', level=2)

doc.add_paragraph(
    'Dashboard admin menampilkan statistik sistem:'
)

dashboard_stats = [
    'Jumlah total user terdaftar',
    'Jumlah role yang tersedia',
    'Jumlah tabel database yang dapat diakses',
]

for stat in dashboard_stats:
    doc.add_paragraph(stat, style='List Bullet')

doc.add_heading('8.2 Manajemen User', level=2)

doc.add_paragraph('Fitur manajemen user mencakup:')

user_features = [
    'Tambah user baru dengan nama, email, password, dan role',
    'Edit user existing (ubah nama, email, role)',
    'Hapus user dari sistem',
    'Search user berdasarkan nama atau email',
    'Filter user berdasarkan role',
    'Pagination (10 user per halaman)',
]

for feature in user_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('8.3 Manajemen Role', level=2)

doc.add_paragraph('Fitur manajemen role:')

role_features = [
    'Tambah role baru dengan nama dan deskripsi',
    'Edit role existing',
    'Hapus role (dapat dihapus, tidak ada role yang dilindungi)',
    'Drag & Drop Permissions:',
]

for feature in role_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph('Fitur Drag & Drop Permissions:', style='List Bullet')

permission_features = [
    'Tabel Tersedia: Tabel yang belum diizinkan untuk role tersebut',
    'Tabel Diizinkan: Tabel yang dapat diakses oleh role',
    'Select All: Pindahkan semua tabel ke "Diizinkan"',
    'Clear All: Pindahkan semua tabel ke "Tersedia"',
    'Indikator Perubahan: Warning jika ada perubahan yang belum disimpan',
]

for feature in permission_features:
    p = doc.add_paragraph(feature, style='List Bullet 2')

doc.add_heading('Cara Kerja Drag & Drop Role', level=3)

drag_drop_steps = [
    'Pilih role dari daftar di sebelah kiri',
    'Drag tabel dari "Tabel Tersedia" ke "Tabel Diizinkan"',
    'Klik "Simpan Akses" untuk menyimpan perubahan',
    'Sistem menampilkan preview perubahan sebelum menyimpan',
    'Border oranye muncul jika ada perubahan belum disimpan',
]

for i, step in enumerate(drag_drop_steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('8.4 Import/Export Excel', level=2)

doc.add_paragraph('Admin panel menyediakan fitur import/export user melalui Excel:')

excel_features = [
    ['Fitur', 'Deskripsi'],
    ['Export Excel', 'Download data semua user ke file Excel (.xlsx)'],
    ['Import Excel', 'Upload file Excel untuk menambah user secara batch'],
    ['Download Template', 'Download template Excel dengan format dan instruksi'],
]

add_formatted_table(excel_features[0], excel_features[1:])

doc.add_heading('Format Template Excel', level=3)

template_columns = [
    'Name: Nama lengkap user',
    'Email: Email user (harus unik)',
    'Password: Password user (akan di-hash)',
    'Role: ID atau nama role',
    'Is Admin: Yes/No (menentukan akses admin)',
]

for col in template_columns:
    doc.add_paragraph(col, style='List Bullet')

doc.add_paragraph(
    'Validasi otomatis saat import akan memastikan:'
)

validations = [
    'Email harus unik (tidak boleh duplikat)',
    'Semua field required (tidak boleh kosong)',
    'Password akan di-hash menggunakan bcrypt',
    'Role harus valid (ada di database)',
]

for val in validations:
    doc.add_paragraph(val, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 9. MCP SERVER INTEGRATION
# ============================================================================
doc.add_heading('9. MCP Server Integration', level=1)

doc.add_heading('9.1 Apa itu MCP?', level=2)

doc.add_paragraph(
    'Model Context Protocol (MCP) adalah protokol standar yang memungkinkan AI model '
    'untuk berinteraksi dengan tools dan data eksternal. Dalam konteks aplikasi ini, MCP '
    'server menyediakan endpoint untuk koneksi dari klien eksternal seperti Claude Desktop '
    'atau AI assistants lainnya.'
)

doc.add_heading('9.2 Konfigurasi MCP', level=2)

doc.add_paragraph(
    'MCP server dikonfigurasi menggunakan package:'
)

mcp_packages = [
    'php-mcp/server: Implementasi MCP server',
    'php-mcp/laravel: Integrasi MCP dengan Laravel',
]

for pkg in mcp_packages:
    doc.add_paragraph(pkg, style='List Bullet')

doc.add_heading('Endpoint MCP', level=3)

doc.add_paragraph(
    'MCP server tersedia di endpoint: GET /mcp'
)

doc.add_paragraph(
    'Endpoint ini dapat diakses oleh klien MCP eksternal untuk berinteraksi dengan '
    'chatbot dan menjalankan tools yang tersedia.'
)

doc.add_heading('9.3 Konfigurasi API Keys', level=2)

doc.add_paragraph(
    'Untuk menggunakan AI models, diperlukan API keys yang dikonfigurasi di .env:'
)

api_keys = [
    ['Variabel', 'Deskripsi'],
    ['OPENROUTER_API_KEY', 'API key untuk OpenRouter (akses berbagai AI models)'],
    ['NVIDIA_API_KEY', 'API key untuk NVIDIA AI services'],
]

add_formatted_table(api_keys[0], api_keys[1:])

doc.add_paragraph(
    'Cara mendapatkan API keys:'
)

api_steps = [
    'OpenRouter: Daftar di https://openrouter.ai dan generate API key',
    'NVIDIA: Daftar di https://build.nvidia.com dan generate API key',
]

for step in api_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 10. OPTIMASI PERFORMA
# ============================================================================
doc.add_heading('10. Optimasi Performa', level=1)

doc.add_paragraph(
    'Aplikasi telah dioptimasi secara menyeluruh untuk memberikan performa terbaik. '
    'Berikut adalah detail optimasi yang telah diimplementasikan:'
)

doc.add_heading('10.1 Database Indexing', level=2)

doc.add_paragraph(
    'Index telah ditambahkan pada kolom-kolom yang sering di-query untuk meningkatkan '
    'kecepatan pencarian:'
)

indexes = [
    ['Tabel', 'Kolom', 'Tipe Index', 'Tujuan'],
    ['chat_messages', 'chat_session_id, created_at', 'Composite', 'Message retrieval cepat'],
    ['chat_messages', 'role', 'Single', 'Filter user/assistant messages'],
    ['chat_sessions', 'user_id, created_at', 'Composite', 'User session listing'],
    ['users', 'role', 'Single', 'RBAC lookups'],
]

add_formatted_table(indexes[0], indexes[1:])

doc.add_paragraph(
    'Expected Impact: 60-80% lebih cepat untuk query database chat history'
).runs[0].font.bold = True

doc.add_heading('10.2 Caching Strategy', level=2)

doc.add_heading('Redis Caching', level=3)

doc.add_paragraph(
    'Redis digunakan untuk caching dengan performa sub-millisecond operations:'
)

redis_config = [
    ['Konfigurasi', 'Value'],
    ['CACHE_STORE', 'redis'],
    ['REDIS_CLIENT', 'phpredis'],
    ['REDIS_HOST', '127.0.0.1'],
    ['REDIS_PORT', '6379'],
    ['SESSION_DRIVER', 'database (bisa diubah ke redis)'],
]

add_formatted_table(redis_config[0], redis_config[1:])

doc.add_heading('Query Result Caching', level=3)

doc.add_paragraph(
    'Hasil query SELECT di-cache selama 60 detik untuk mengurangi beban database:'
)

query_cache_details = [
    'Cache key dihasilkan dari SQL hash + user ID',
    'TTL (Time To Live): 60 detik',
    'Otomatis cache hasil query identik',
    'User-specific caching untuk keamanan data',
]

for detail in query_cache_details:
    doc.add_paragraph(detail, style='List Bullet')

doc.add_heading('Configuration & Route Caching', level=3)

doc.add_paragraph(
    'Di production, config dan routes di-cache untuk mengurangi bootstrap time:'
)

cache_commands = [
    'php artisan config:cache - Cache konfigurasi',
    'php artisan route:cache - Cache routing',
    'php artisan view:cache - Cache view templates',
]

for cmd in cache_commands:
    doc.add_paragraph(cmd, style='List Bullet')

doc.add_heading('Expected Impact dari Caching', level=3)

cache_impacts = [
    ['Metrik', 'Sebelum', 'Setelah', 'Improvement'],
    ['Cache operations', '10-100ms (DB)', '< 1ms (Redis)', '90-99% lebih cepat'],
    ['Repeated queries', 'Hit database', 'Dari cache', '30-40% lebih cepat'],
    ['Bootstrap time', '150-300ms', '50-100ms', '60-70% lebih cepat'],
]

add_formatted_table(cache_impacts[0], cache_impacts[1:])

doc.add_heading('10.3 Laravel Octane', level=2)

doc.add_paragraph(
    'Laravel Octane meningkatkan performa dengan menjaga aplikasi tetap di memori:'
)

doc.add_heading('Apa itu Laravel Octane?', level=3)

doc.add_paragraph(
    'Octaneboot aplikasi sekali dan menjaga agar tetap di memori. Request selanjutnya '
    'ditangani oleh aplikasi yang sudah di-boot, menghasilkan response time 5-10x lebih cepat.'
)

doc.add_heading('Performance Comparison', level=3)

octane_comparison = [
    ['Metrik', 'PHP-FPM', 'Octane (Swoole)', 'Improvement'],
    ['Response time', '150-300ms', '30-60ms', '5-7x lebih cepat'],
    ['Requests/second', '100-200', '1000-2000+', '10x lebih banyak'],
    ['Concurrent connections', '10-20', '500-1000+', '50x lebih banyak'],
    ['Memory per request', '40-80MB', '5-10MB', '80% reduction'],
]

add_formatted_table(octane_comparison[0], octane_comparison[1:])

doc.add_heading('Install Octane', level=3)

octane_steps = [
    'composer require laravel/octane',
    'php artisan octane:install --server=swoole',
    'php artisan octane:start',
]

for step in octane_steps:
    doc.add_paragraph(step, style='List Bullet')

doc.add_paragraph(
    'Catatan: Swoole tidak support Windows native. Gunakan Docker atau WSL2 untuk Windows.'
)

doc.add_heading('10.4 Chat Pagination', level=2)

doc.add_paragraph(
    'Chat history dimuat dengan pagination untuk performa optimal:'
)

pagination_details = [
    'Default: 50 pesan terakhir dimuat',
    'Cursor-based pagination untuk performa',
    'Tombol "Muat Pesan Lebih Awal" untuk load older messages',
    'Maintain scroll position saat load older messages',
    'Expected Impact: 70% reduction in initial load time',
]

for detail in pagination_details:
    doc.add_paragraph(detail, style='List Bullet')

doc.add_heading('10.5 Combined Performance Impact', level=2)

doc.add_paragraph(
    'Setelah semua optimasi diimplementasikan:'
)

combined_impact = [
    ['Metrik', 'Sebelum', 'Sesudah', 'Improvement'],
    ['Chat history load', '3-5 detik', '200-400ms', '90% lebih cepat'],
    ['First chatbot response', '500-1000ms', '150-300ms', '60-70% lebih cepat'],
    ['Database queries/chat', '20-50', '5-15', '70-80% reduction'],
    ['Concurrent users', '20-50', '200-500+', '10x capacity'],
    ['Memory per request', '80-120MB', '20-40MB', '70% reduction'],
    ['Requests/second', '50-100', '500-1000+', '10x throughput'],
]

add_formatted_table(combined_impact[0], combined_impact[1:])

doc.add_page_break()

# ============================================================================
# 11. DEPLOYMENT
# ============================================================================
doc.add_heading('11. Deployment', level=1)

doc.add_heading('11.1 Deployment di Ubuntu 22.04', level=2)

doc.add_heading('Prerequisites', level=3)

prereqs = [
    '✅ Ubuntu 22.04 server',
    '✅ PHP 8.2-FPM installed',
    '✅ Nginx installed',
    '✅ PostgreSQL installed',
    '✅ Laravel project sudah di-upload ke server',
]

for prereq in prereqs:
    doc.add_paragraph(prereq, style='List Bullet')

doc.add_heading('Yang Akan Diinstall', level=3)

to_install = [
    'Redis server',
    'PHP Redis extension',
    'Laravel Octane + Swoole (optional tapi recommended)',
    'Optimasi Nginx + PHP-FPM',
]

for item in to_install:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Langkah Deployment', level=3)

deployment_steps = [
    ('Setup Redis Server', [
        'sudo apt update',
        'sudo apt install redis-server -y',
        'sudo systemctl enable redis-server',
        'sudo systemctl start redis-server',
        'redis-cli ping  # Harusnya: PONG'
    ]),
    ('Update Laravel .env', [
        'CACHE_STORE=redis',
        'APP_ENV=production',
        'APP_DEBUG=false',
        'LOG_LEVEL=error'
    ]),
    ('Set File Permissions', [
        'sudo chown -R www-data:www-data /path/to/project',
        'sudo find /path/to/project -type f -exec chmod 644 {} \\;',
        'sudo find /path/to/project -type d -exec chmod 755 {} \\;',
        'sudo chmod -R 775 /path/to/project/storage',
        'sudo chmod -R 775 /path/to/project/bootstrap/cache'
    ]),
    ('Install Dependencies & Build', [
        'composer install --optimize-autoloader --no-dev',
        'npm ci',
        'npm run build'
    ]),
    ('Run Migrations & Clear Caches', [
        'php artisan migrate --force',
        'php artisan config:clear',
        'php artisan cache:clear',
        'php artisan view:clear'
    ]),
    ('Cache Configuration (Production Mode)', [
        'php artisan config:cache',
        'php artisan route:cache',
        'php artisan view:cache',
        'composer dump-autoload --optimize'
    ]),
    ('Configure Nginx', [
        'Buat virtual host di /etc/nginx/sites-available/',
        'Setup server block dengan root ke public/',
        'Config PHP-FPM socket',
        'Config SSE untuk chatbot streaming',
        'sudo nginx -t && sudo systemctl reload nginx'
    ]),
    ('Optimize PHP-FPM', [
        'pm.max_children = 50',
        'pm.start_servers = 5',
        'pm.max_requests = 500',
        'Enable OPcache dengan memory 256MB',
        'sudo systemctl restart php8.2-fpm'
    ]),
]

for i, (step, commands) in enumerate(deployment_steps, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'Step {i}: {step}')
    run.font.bold = True
    
    for cmd in commands:
        doc.add_paragraph(cmd).style = 'No Spacing'
    doc.add_paragraph()

doc.add_heading('11.2 Production Checklist', level=2)

doc.add_heading('Pre-Deployment', level=3)

pre_deploy = [
    '☐ Update .env dengan konfigurasi production',
    '☐ Set APP_ENV=production',
    '☐ Set APP_DEBUG=false',
    '☐ Konfigurasi database PostgreSQL',
    '☐ Setup Redis server',
    '☐ Install PHP extensions required',
]

for item in pre_deploy:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Deployment', level=3)

deploy_checklist = [
    '☐ Upload code ke server (git pull / rsync)',
    '☐ composer install --optimize-autoloader --no-dev',
    '☐ npm ci && npm run build',
    '☐ php artisan migrate --force',
    '☐ php artisan config:cache',
    '☐ php artisan route:cache',
    '☐ php artisan view:cache',
    '☐ Set file permissions (www-data)',
    '☐ Restart PHP-FPM',
    '☐ Restart Nginx',
]

for item in deploy_checklist:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Post-Deployment', level=3)

post_deploy = [
    '☐ Test website accessibility (curl -I http://domain)',
    '☐ Test login functionality',
    '☐ Test chatbot response',
    '☐ Verify Redis connection',
    '☐ Check database indexes',
    '☐ Review error logs (storage/logs/laravel.log)',
    '☐ Monitor Nginx error log',
    '☐ Benchmark performance (ab -n 100 -c 10)',
]

for item in post_deploy:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Security Hardening', level=3)

security_items = [
    '☐ Setup firewall (UFW): allow SSH, HTTP, HTTPS',
    '☐ Setup SSL (Let\'s Encrypt/Certbot)',
    '☐ Secure .env file (chmod 640)',
    '☐ Deny access to hidden files di Nginx',
    '☐ Disable PHP execution di storage/',
    '☐ Setup log rotation',
    '☐ Enable rate limiting (optional)',
]

for item in security_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 12. TROUBLESHOOTING
# ============================================================================
doc.add_heading('12. Troubleshooting', level=1)

doc.add_heading('12.1 Database Issues', level=2)

doc.add_heading('Error: permission denied for table', level=3)
doc.add_paragraph('Solusi:')
doc.add_paragraph('GRANT ALL PRIVILEGES ON ALL TABLES IN SCHEMA public TO postgres;', style='No Spacing')
doc.add_paragraph('GRANT ALL PRIVILEGES ON ALL SEQUENCES IN SCHEMA public TO postgres;', style='No Spacing')

doc.add_heading('Error: SQLSTATE[08006] - Connection refused', level=3)
doc.add_paragraph('Solusi:')
connection_fix = [
    'Pastikan PostgreSQL service berjalan: sudo systemctl status postgresql',
    'Cek konfigurasi .env (host, port, username, password)',
    'Pastikan PostgreSQL menerima koneksi TCP',
    'Edit postgresql.conf: listen_addresses = \'localhost\'',
]
for fix in connection_fix:
    doc.add_paragraph(fix, style='List Bullet')

doc.add_heading('12.2 Application Issues', level=2)

doc.add_heading('Error: class not found setelah composer install', level=3)
doc.add_paragraph('Solusi:')
doc.add_paragraph('composer dump-autoload', style='No Spacing')

doc.add_heading('Error: npm run build gagal', level=3)
doc.add_paragraph('Solusi:')
doc.add_paragraph('rm -rf node_modules package-lock.json', style='No Spacing')
doc.add_paragraph('npm install', style='No Spacing')
doc.add_paragraph('npm run build', style='No Spacing')

doc.add_heading('12.3 Redis Issues', level=2)

doc.add_heading('Connection Refused Error', level=3)
redis_fix = [
    'Pastikan Redis server running: redis-cli ping (harusnya PONG)',
    'Cek firewall settings allow port 6379',
    'Verifikasi REDIS_HOST dan REDIS_PORT di .env',
]
for fix in redis_fix:
    doc.add_paragraph(fix, style='List Bullet')

doc.add_heading('PHP Extension Not Found', level=3)
doc.add_paragraph('Solusi:')
doc.add_paragraph('Ganti REDIS_CLIENT=phpredis ke REDIS_CLIENT=predis di .env', style='No Spacing')
doc.add_paragraph('composer require predis/predis', style='No Spacing')

doc.add_heading('12.4 Crawler Issues', level=2)

doc.add_heading('Crawler Gagal Login', level=3)
crawler_fix = [
    'Cek credentials di app/Console/Commands/CrawlDocumentation.php',
    'Pastikan website ERP Guidance dapat diakses',
    'Cek cookie/session handling di server',
]
for fix in crawler_fix:
    doc.add_paragraph(fix, style='List Bullet')

doc.add_heading('12.5 Queue Issues', level=2)

doc.add_heading('Queue Worker Tidak Berjalan', level=3)
queue_fix = [
    'Cek status supervisor: sudo supervisorctl status',
    'Restart worker: sudo supervisorctl restart mcp-worker:*',
    'Cek log: tail -f storage/logs/laravel.log',
]
for fix in queue_fix:
    doc.add_paragraph(fix, style='List Bullet')

doc.add_heading('12.6 Session Issues', level=2)

doc.add_heading('Session Error di Production', level=3)
session_fix = [
    'php artisan config:clear',
    'php artisan cache:clear',
    'php artisan session:table',
    'php artisan migrate',
]
for fix in session_fix:
    doc.add_paragraph(fix, style='List Bullet')

doc.add_page_break()

# ============================================================================
# 13. MAINTENANCE AND MONITORING
# ============================================================================
doc.add_heading('13. Maintenance dan Monitoring', level=1)

doc.add_heading('13.1 Daily Maintenance', level=2)

daily_tasks = [
    ['Task', 'Command', 'Deskripsi'],
    ['Monitor error logs', 'tail -f storage/logs/laravel.log', 'Cek error yang terjadi'],
    ['Check Redis memory', 'redis-cli INFO memory', 'Monitor Redis usage'],
    ['Check disk space', 'df -h', 'Pastikan storage cukup'],
]

add_formatted_table(daily_tasks[0], daily_tasks[1:])

doc.add_heading('13.2 Weekly Maintenance', level=2)

weekly_tasks = [
    ['Task', 'Command', 'Deskripsi'],
    ['Review cache hit rates', 'redis-cli INFO stats', 'Efektivitas caching'],
    ['Check query performance', 'Monitor slow queries', 'Optimize jika perlu'],
    ['Monitor queue workers', 'sudo supervisorctl status', 'Pastikan worker running'],
    ['Review error logs', 'cat storage/logs/laravel.log', 'Analisis pola error'],
]

add_formatted_table(weekly_tasks[0], weekly_tasks[1:])

doc.add_heading('13.3 Monthly Maintenance', level=2)

monthly_tasks = [
    ['Task', 'Command', 'Deskripsi'],
    ['Update dependencies', 'composer update && npm update', 'Update packages'],
    ['Review slow queries', 'pg_stat_statements', 'Identifikasi query lambat'],
    ['Test rollback procedures', 'php artisan migrate:rollback', 'Validasi backup'],
    ['Database backup', 'pg_dump database > backup.sql', 'Backup rutin'],
    ['Clean old logs', 'truncate storage/logs/laravel.log', 'Bersihkan logs'],
]

add_formatted_table(monthly_tasks[0], monthly_tasks[1:])

doc.add_heading('13.4 Monitoring Commands', level=2)

doc.add_heading('Application Monitoring', level=3)

monitoring_cmds = [
    ['Task', 'Command'],
    ['Check application status', 'php artisan status'],
    ['Clear all caches', 'php artisan optimize:clear'],
    ['View routes', 'php artisan route:list'],
    ['Check config', 'php artisan config:show'],
    ['Database migrations status', 'php artisan migrate:status'],
    ['Queue failed jobs', 'php artisan queue:failed'],
]

add_formatted_table(monitoring_cmds[0], monitoring_cmds[1:])

doc.add_heading('Database Monitoring', level=3)

db_monitoring = [
    ['Task', 'Command'],
    ['Check table sizes', '\\dt+'],
    ['Check indexes', '\\di'],
    ['View active connections', 'SELECT * FROM pg_stat_activity;'],
    ['Table statistics', 'SELECT * FROM pg_stat_user_tables;'],
    ['Index usage', 'SELECT * FROM pg_stat_user_indexes;'],
]

add_formatted_table(db_monitoring[0], db_monitoring[1:])

doc.add_heading('Server Monitoring', level=3)

server_monitoring = [
    ['Task', 'Command'],
    ['CPU usage', 'top atau htop'],
    ['Memory usage', 'free -h'],
    ['Disk usage', 'df -h'],
    ['Nginx status', 'sudo systemctl status nginx'],
    ['PHP-FPM status', 'sudo systemctl status php8.2-fpm'],
    ['Redis status', 'sudo systemctl status redis-server'],
    ['PostgreSQL status', 'sudo systemctl status postgresql'],
]

add_formatted_table(server_monitoring[0], server_monitoring[1:])

doc.add_heading('13.5 Backup Strategy', level=2)

doc.add_heading('Database Backup', level=3)

backup_commands = [
    ['Type', 'Command'],
    ['Full backup', 'pg_dump db_penjualan > backup_$(date +%Y%m%d).sql'],
    ['Compressed backup', 'pg_dump db_penjualan | gzip > backup_$(date +%Y%m%d).sql.gz'],
    ['Restore backup', 'psql db_penjualan < backup_20260409.sql'],
]

add_formatted_table(backup_commands[0], backup_commands[1:])

doc.add_heading('File Backup', level=3)

file_backup_items = [
    '.env file (konfigurasi sensitif)',
    'storage/app/ (user uploads)',
    'storage/logs/ (log files)',
]

for item in file_backup_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('13.6 Log Management', level=2)

doc.add_heading('Log Files Location', level=3)

log_files = [
    ['Log File', 'Location', 'Purpose'],
    ['Laravel Log', 'storage/logs/laravel.log', 'Application errors & events'],
    ['Nginx Access', '/var/log/nginx/access.log', 'HTTP requests'],
    ['Nginx Error', '/var/log/nginx/error.log', 'Nginx errors'],
    ['PHP-FPM Log', '/var/log/php8.2-fpm.log', 'PHP-FPM errors'],
    ['PostgreSQL Log', '/var/log/postgresql/', 'Database logs'],
    ['Redis Log', '/var/log/redis/', 'Redis logs'],
]

add_formatted_table(log_files[0], log_files[1:])

doc.add_heading('Log Rotation Setup', level=3)

doc.add_paragraph('Buat file /etc/logrotate.d/laravel:')
doc.add_paragraph(
    '/path/to/project/storage/logs/*.log {\n'
    '    daily\n'
    '    missingok\n'
    '    rotate 14\n'
    '    compress\n'
    '    delaycompress\n'
    '    notifempty\n'
    '    create 0640 www-data www-data\n'
    '    sharedscripts\n'
    '    postrotate\n'
    '        systemctl reload php8.2-fpm > /dev/null 2>&1 || true\n'
    '    endscript\n'
    '}',
    style='No Spacing'
)

doc.add_page_break()

# ============================================================================
# CLOSING
# ============================================================================
doc.add_heading('Penutup', level=1)

doc.add_paragraph(
    'Dokumentasi ini berisi panduan lengkap untuk aplikasi MCP PostgreSQL Chatbot & ERP Integration. '
    'Semua aspek telah dibahas mulai dari instalasi, konfigurasi, fitur, deployment, hingga maintenance.'
)

doc.add_paragraph(
    'Untuk pertanyaan lebih lanjut atau bantuan, silakan menghubungi tim development atau '
    'merujuk pada README.md dan dokumen panduan lainnya yang tersedia di repository:'
)

additional_docs = [
    'REDIS_SETUP.md - Panduan setup Redis',
    'QUEUE_SETUP.md - Panduan setup queue system',
    'OCTANE_SETUP.md - Panduan setup Laravel Octane',
    'PRODUCTION_CACHING.md - Panduan caching production',
    'DEPLOY_UBUNTU.md - Panduan deploy di Ubuntu',
    'PERFORMANCE_OPTIMIZATION_SUMMARY.md - Ringkasan optimasi performa',
]

for doc_item in additional_docs:
    doc.add_paragraph(doc_item, style='List Bullet')

doc.add_paragraph()

# Footer
footer_para = doc.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run('=' * 60)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

footer_para2 = doc.add_paragraph()
footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para2.add_run('Akhir Dokumentasi')
run.font.size = Pt(14)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

footer_para3 = doc.add_paragraph()
footer_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para3.add_run(f'Dibuat pada: {datetime.datetime.now().strftime("%d %B %Y")}')
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = r'D:\MCP Versi Web\mcp-postgresql\DOKUMENTASI_MCP_CHATBOT.docx'
doc.save(output_path)

print(f'Dokumentasi berhasil dibuat di: {output_path}')
